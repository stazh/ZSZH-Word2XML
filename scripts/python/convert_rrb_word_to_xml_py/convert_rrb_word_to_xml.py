from docx import Document # type: ignore
from docx.document import Document as _Document # type: ignore
from docx.oxml.text.paragraph import CT_P # type: ignore
from docx.oxml.table import CT_Tbl # type: ignore
from docx.table import _Cell, Table, _Row # type: ignore
from docx.text.paragraph import Paragraph # type: ignore
import xlsxwriter # type: ignore
import re
import pandas as pd
import argparse
import os
import shutil
from os import walk
import zipfile
from bs4 import BeautifulSoup, NavigableString

ERRORFILE = xlsxwriter.Workbook("errorfile.xlsx")
ERRORWS = ERRORFILE.add_worksheet()
ERRORCOUNTER = 0
ERROR_FOUND = False

def create_doc(name):
    return Document(name)

def load_table(metadata_file):
    xls = pd.ExcelFile(metadata_file)
    df = xls.parse(xls.sheet_names[0])
    dict_new = dict(zip(df.Signatur, df.ID))
    return dict_new

def create_scopeLink(signatur, database,docname):
    signaturber = signatur.replace(" StAZH ","")
    signaturber = signaturber.replace("StAZH ","")
    signaturber = signaturber.replace(" StAZH","")
    signaturber = signaturber.replace("StAZH","")
    signaturber = signaturber.replace("\n","")
    signaturber = signaturber.replace("–"," - ")
    signaturber = signaturber.replace("  "," ")
    signaturber = signaturber.strip()

    try:
        id = database[signaturber]
    except:
        id = 1.0
        log_error("Scopelink nicht in Datenbank: "+signaturber,docname)
    if isinstance(id,float) or isinstance(id,str) or isinstance(id,int):
        return "https://suche.staatsarchiv.djiktzh.ch/detail.aspx?ID=" + str(int(id))
    else:
        log_error("Scopelink nicht in Datenbank: "+signatur,docname)
        return "https://suche.staatsarchiv.djiktzh.ch/detail.aspx?ID="


def word_header_df(doc):
    table = doc.tables[0]
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        data.append(row_data)
    return pd.DataFrame(data)

def check_table(name: str, doc):
    tablelist = []
    for table in doc.tables:
        tablelist.append(table)
    return tablelist

def log_error(error:str,filename:str):
    global ERRORCOUNTER
    global ERRORWS
    global ERRORFILE
    global ERROR_FOUND
    ERROR_FOUND = True
    ERRORCOUNTER = ERRORCOUNTER + 1
    ERRORWS.write('A'+str(ERRORCOUNTER), filename)
    ERRORWS.write('B'+str(ERRORCOUNTER), error)

def get_transcriptation_date(doc,docname):
    #TODO: Errrorhandling
    transkriberfound = False
    trans_date = ""
    for para in doc.paragraphs:  
        if "[Transkript:" in str(para.text):
            try:
                footer = str(para.text)
                footer_list = footer.split(":")
                transcriptor_and_date = footer_list[2].split("/")
                trans_date = transcriptor_and_date[1].strip()
                trans_date = trans_date.replace("]","").strip()
                transkriberfound = True
            except:
                log_error("Etwas am Ende des Dokuments (Transkriptor) stimmt nicht.",docname)
    if not transkriberfound:
        log_error("Keinen Transkriptor am Dokumentende gefunden",docname)
    return trans_date


def date_adjuster(datestr,docname):
    splitlist = datestr.split(".")
    try:
        if len(splitlist) == 3:
            if len(splitlist[0]) == 1:
               splitlist[0] = "0" + splitlist[0]
            if len(splitlist[1]) == 1:
               splitlist[1] = "0" + splitlist[1]
            return splitlist[2] + "-" + splitlist[1] + "-" + splitlist[0]
        elif len(splitlist) == 2:
            if len(splitlist[0]) == 1:
               splitlist[0] = "0" + splitlist[0]
            return splitlist[1] + "-" + splitlist[0]
        elif len(splitlist) == 1:
            return splitlist[0]
        else:
            log_error("Datum nicht im üblichen Format: "+ datestr,docname)
            return datestr
    except:
        log_error("Datum nicht im üblichen Format: "+ datestr,docname)
        return datestr

def parse_ident(name: str):
    tf = word_header_df(0, name)
    tf = tf.set_index(0)
    sig = tf.loc["Signatur"].values[0]
    x = sig.split()
    retident = ""
    for i in range(1, len(x)):
        retident += x[i] + " "
    return retident[:-1]

def is_one_date(date):
    if "–" in date:
        return False
    else:
        return True

def graphic_url_creator(file_name):
    file_name = file_name.replace("/Users/stazh/Documents/GithubRepos/ZSZHWordToXMLConverter/","")
    file_name = file_name.replace("_t","_p.pdf")
    return file_name

def iter_physical_row_cells(table, row):
    tr = row._tr
    for tc in tr.tc_lst:
        yield _Cell(tc, table)

def doc_name_withoutdoc(name: str):
    parse_ident(name).replace(" ", "_").replace(".", "_").replace("/", "_")

def extract_pictures(name):
    archive = zipfile.ZipFile(name)
    picturelist = []
    for file in archive.filelist:
        if file.filename.startswith("word/media/"):
            number += 1
            picture_name = doc_name_withoutdoc + "Grafik" + number + ".jpg"
            picturelist.append(picture_name)
            archive.extract(file, picture_name)
    return picturelist

def get_title(metadata,word_file_with_path):
    try:
        return metadata.loc["Titel"].values[0]
    except:
        log_error("Kein Titel in Metadaten",word_file_with_path)
        return ""
    
def get_signatur(metadata,word_file_with_path):
    try:
        return metadata.loc["Signatur"].values[0]
    except:
        log_error("Keine Signatur in Metadaten",word_file_with_path)
        return ""
    
def get_idno(metadata,word_file_with_path):
    try:
        return metadata.loc["P."].values[0]
    except:
        log_error("Kein P. in Metadaten",word_file_with_path)
        return ""

def get_date(metadata,word_file_with_path):
    try:
        return metadata.loc["Datum"].values[0]
    except:
        log_error("Kein Datum in Metadaten",word_file_with_path)
        return ""

def xml_outputer(word_file, word_file_with_path,doc, metadata, database,subdir):
    global ERROR_FOUND 
    ERROR_FOUND = False
    subdirfehler = subdir +"_Fehler"
    subdirxml = subdir + "_XML"
    if not os.path.exists(subdirxml):
        os.makedirs(subdirxml)
    word_file_fehler_with_path = subdirfehler + "/" + word_file
    xml_file_with_path = subdirxml + "/" + word_file
    xml_file_with_path = xml_file_with_path.replace(".docx",".xml")
    f = open(xml_file_with_path, "a+", encoding="utf-8")
    soup = BeautifulSoup(f,features='xml')
    prolog = BeautifulSoup('<?xml-stylesheet type="text/xsl" href="../../Ressourcen/Stylesheet.xsl"?>',features='xml')
    soup.append(prolog)
    metadata = metadata.set_index(0)
    date = get_date(metadata,word_file_with_path).strip()
    isOneDate = is_one_date(date)
    firstdate = ""
    seconddate = ""
    if isOneDate:
        firstdate = date
    else:
        datelist = date.split("–")
        firstdate = datelist[0].strip()
        seconddate = datelist[1].strip()
    TEI = soup.new_tag("TEI")
    TEI['xmlns'] = "http://www.tei-c.org/ns/1.0"
    TEI['xmlns:xsi'] = "http://www.w3.org/2001/XMLSchema-instance"
    soup.append(TEI)
    teiHeader = soup.new_tag("teiHeader")
    TEI.append(teiHeader)
    fileDesc = soup.new_tag("fileDesc")
    teiHeader.append(fileDesc)
    titleStmt = soup.new_tag("titleStmt")
    fileDesc.append(titleStmt)
    titleHeader = soup.new_tag("title")
    title = get_title(metadata,word_file_with_path)
    titleHeader.string = title
    titleStmt.append(titleHeader)
    respStmt = soup.new_tag("respStmt")
    titleStmt.append(respStmt)
    transcriptorin_kuerzel = "#OCR/Team TKR"
    resp = soup.new_tag("resp", ref=transcriptorin_kuerzel, key="transcript")
    respStmt.append(resp)
    resp.string = "Transkript: " + transcriptorin_kuerzel.replace("#","")
    transcript_date = get_transcriptation_date(doc, word_file_with_path)
    respdate = soup.new_tag("date", when=date_adjuster(transcript_date,word_file_with_path))
    respdate.string = transcript_date
    resp.append(respdate)
    respStmtName = soup.new_tag("name", key="editor")
    respStmtName.string = "Staatsarchiv des Kantons Zürich"
    respStmt.append(respStmtName)
    publicationStmt = soup.new_tag("publicationStmt")
    fileDesc.append(publicationStmt)
    authority = soup.new_tag("authority")
    authority.string = "Staatsarchiv des Kantons Zürich"
    publicationStmt.append(authority)
    publicationStmtDate = soup.new_tag("date")
    section = doc.sections[0]
    footer = section.footer
    footerpara = footer.paragraphs[0]
    publicationStmtDate.string = footerpara.text[-4:]
    publicationStmt.append(publicationStmtDate)
    pubPlace = soup.new_tag("pubPlace")
    pubPlace.string = "Zürich"
    seriesStmt = soup.new_tag("seriesStmt")
    fileDesc.append(seriesStmt)
    seriesStmtTitle = soup.new_tag("title")
    seriesStmtTitle.string = "Regierungsratsbeschlüsse seit 1803 online"
    seriesStmt.append(seriesStmtTitle)
    editor = soup.new_tag("editor")
    editor.string = "Staatsarchiv des Kantons Zürich"
    seriesStmt.append(editor)
    sourceDesc = soup.new_tag("sourceDesc")
    fileDesc.append(sourceDesc)
    bibl = soup.new_tag("bibl")
    sourceDesc.append(bibl)
    biblTitle = soup.new_tag("title")
    biblTitle.string = title
    bibl.append(biblTitle)
    if isOneDate:
        biblDate = soup.new_tag("date",when=date_adjuster(firstdate,word_file_with_path))
        biblDate.string = firstdate
    else:
        biblDate = soup.new_tag("date")
        biblDate['from'] = date_adjuster(firstdate,word_file_with_path)
        biblDate['to'] = date_adjuster(seconddate,word_file_with_path)
        biblDate.string = firstdate + "–" + seconddate
    bibl.append(biblDate)
    ident = soup.new_tag("ident")
    signatur = get_signatur(metadata, word_file_with_path)
    ident.string = signatur
    bibl.append(ident)
    idno = soup.new_tag("idno")
    idno.string = get_idno(metadata, word_file_with_path)
    bibl.append(idno)
    edition = soup.new_tag("edition")
    # important to change if other document
    edition.string = "Regierungsratbeschlüsse seit 1803 online"
    bibl.append(edition)
    figure = soup.new_tag("figure")
    graphic = soup.new_tag("graphic",url=graphic_url_creator(word_file_with_path.replace(subdir,"").replace(".docx","")))
    figure.append(graphic)
    bibl.append(figure)
    ref = soup.new_tag("ref", target=create_scopeLink(signatur,database,word_file_with_path))
    bibl.append(ref)
    TEItext = soup.new_tag("text")
    TEI.append(TEItext)
    body = soup.new_tag("body")
    TEItext.append(body)
    tablecounter = 0
    for block in iter_doc_blockitems(doc):
        # always skip first table since we handle it seperatly
        if isinstance(block, Table):
            # always skip first table because in header
            if tablecounter == 0:
                tablecounter += 1
                continue
            childp = soup.new_tag("p")
            body.insert(len(body),childp)
            childtable = soup.new_tag("table")
            childp.append(childtable)
            for row in block.table.rows:
                cells = list(iter_physical_row_cells(block.table, row))
                childrow = soup.new_tag("row")
                childtable.append(childrow)
                for cell in cells:
                    childcell = soup.new_tag("cell")
                    childrow.append(childcell)
                    if not cell.text == '':
                        pagebreaks = re.findall(r'\[p\.\s.*?\]', cell.text)
                        for pagebreak in pagebreaks:
                            pagebreaktemp = pagebreak[4:]
                            pagebreaklist = pagebreaktemp.split("]")
                            pagenumber = pagebreaklist[0]
                            cell.text = cell.text.replace(" "+pagebreak+" ","STARTPB"+pagenumber+"ENDPB")
                            cell.text = cell.text.replace(" "+pagebreak,"STARTPB"+pagenumber+"ENDPB")
                            cell.text = cell.text.replace(pagebreak+" ","STARTPB"+pagenumber+"ENDPB")
                            cell.text = cell.text.replace(pagebreak,"STARTPB"+pagenumber+"ENDPB")
                            cell.text = cell.text.replace("[up.]","STARTPBupENDPB")
                            cell.text = cell.text.replace(" // ","")
                            cell.text = cell.text.replace(" //","")
                            cell.text = cell.text.replace("// ","")
                            cell.text = cell.text.replace("//","")
                    if "STARTPB" in cell.text:
                        textbeforePB = cell.text[:cell.text.find("STARTPB")]
                        childcell.insert(len(childcell), NavigableString(textbeforePB))
                        templist = cell.text.split("STARTPB")
                        pagenumber = templist[1][:templist[1].find("ENDPB")]
                        pb = soup.new_tag("pb",n=pagenumber)
                        childcell.insert(len(childcell),pb)
                        templist = cell.text.split("ENDPB")
                        childcell.insert(len(childcell), NavigableString(templist[1]))
                    else:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.italic:
                                    childitalic = soup.new_tag("hi",rend="italic")
                                    childitalic.string = run.text
                                    childcell.insert(len(childcell),childitalic)
                                elif run.underline:
                                    childunderline = soup.new_tag("hi",rend="underline")
                                    childunderline.string = run.text
                                    childcell.insert(len(childcell),childunderline)
                                elif run.font.subscript:
                                    childsub = soup.new_tag("hi",rend="sub")
                                    childsub.string = run.text
                                    childcell.insert(len(childcell),childsub)
                                elif run.font.superscript:
                                    childsup = soup.new_tag("hi",rend="sup")
                                    childsup.string = run.text
                                    childcell.insert(len(childcell),childsup)
                                else:
                                    childcell.insert(len(childcell),NavigableString(run.text))

        elif isinstance(block, Paragraph):
            if not block.text == '':
                pagebreaks = re.findall(r'\[p\.\s.*?\]', block.text)
                for pagebreak in pagebreaks:
                    pagebreaktemp = pagebreak[4:]
                    pagebreaklist = pagebreaktemp.split("]")
                    pagenumber = pagebreaklist[0]
                    block.text = block.text.replace(" "+pagebreak+" ","STARTPB"+pagenumber+"ENDPB")
                    block.text = block.text.replace(" "+pagebreak,"STARTPB"+pagenumber+"ENDPB")
                    block.text = block.text.replace(pagebreak+" ","STARTPB"+pagenumber+"ENDPB")
                    block.text = block.text.replace(pagebreak,"STARTPB"+pagenumber+"ENDPB")
                    block.text = block.text.replace("[up.]","STARTPBupENDPB")
                    block.text = block.text.replace(" // ","")
                    block.text = block.text.replace(" //","")
                    block.text = block.text.replace("// ","")
                    block.text = block.text.replace("//","")
            childp = soup.new_tag("p")          
            if "[Transkript:" not in block.text:
                childp = soup.new_tag("p")
                for run in block.runs:
                    if "STARTPB" in run.text:
                        templist = run.text.split("STARTPB")
                        if len(templist) > 1:
                            counter = 0
                            for elem in templist:
                                if counter == 0:
                                    textbeforePB = templist[0]
                                    childp.insert(len(childp), NavigableString(textbeforePB))
                                else:
                                    templist = elem.split("ENDPB")
                                    if len(templist) == 2:
                                        pagenumber = templist[0]
                                        pb = soup.new_tag("pb",n=pagenumber)
                                        childp.insert(len(childp),pb)
                                        childp.insert(len(childp), NavigableString(templist[1]))
                                    else:                             
                                        childp.insert(len(childp),NavigableString(''.join(templist)))
                                        log_error("Irgendetwas stimmt mit Pagesbreaks nicht.", word_file_with_path)
                                counter+=1
                        else:
                            childp.insert(len(childp),NavigableString(run.text))
                            log_error("Irgendetwas stimmt mit den Pagesbreaks nicht.", word_file_with_path)
                    elif run.italic:
                        childitalic = soup.new_tag("hi",rend="italic")
                        childitalic.string = run.text
                        childp.insert(len(childp),childitalic)
                    elif run.underline:
                        childunderline = soup.new_tag("hi",rend="underline")
                        childunderline.string = run.text
                        childp.insert(len(childp),childunderline)
                    elif run.font.subscript:
                        childsub = soup.new_tag("hi",rend="sub")
                        childsub.string = run.text
                        childp.insert(len(childp),childsub)
                    elif run.font.superscript:
                        childsup = soup.new_tag("hi",rend="sup")
                        childsup.string = run.text
                        childp.insert(len(childp),childsup)
                    else:
                        childp.insert(len(childp),NavigableString(run.text))
            body.insert(len(body),childp)
    if "[p." in soup.text:
        log_error("Immernoch ein [p. im Text", word_file_with_path)
    if ERROR_FOUND:
        if not os.path.exists(subdirfehler):
            os.makedirs(subdirfehler)
        shutil.copyfile(word_file_with_path,word_file_fehler_with_path)
    else:
        f.write(str(soup))
        f.close()


def iter_doc_blockitems(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("Something is not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def main():
     # Argumentparser einrichten
    parser = argparse.ArgumentParser(description="Verarbeitet einen Ordner mit Word-Dateien und eine Excel-Datei mit Metadaten.")
    parser.add_argument('input_folder', type=str, help="Pfad zum Ordner mit Word-Dateien")
    parser.add_argument('metadata_file', type=str, help="Pfad zur Excel-Datei mit Metadaten")

    args = parser.parse_args()
    input_folder = args.input_folder
    metadata_file = args.metadata_file

    # Überprüfen, ob der Ordner existiert
    if not os.path.isdir(input_folder):
        print(f"Der angegebene Ordner '{input_folder}' existiert nicht.")
        return

    # Überprüfen, ob die Excel-Datei existiert und geladen werden kann
    if not os.path.isfile(metadata_file) or not metadata_file.endswith(('.xlsx', '.xls')):
        print(f"Die angegebene Excel-Datei '{metadata_file}' ist ungültig oder existiert nicht.")
        return

    # word_path = "/Users/stazh/Documents/GithubRepos/ZSZHWordToXMLConverter/TKR_RRB_Transkripte"
    print("Loading data...")
    database = load_table(metadata_file)
    print("Starting with convertion...")
    for subdir, dirs, files in walk(input_folder):
        for file in files:
            file_with_path = os.path.join(subdir, file)
            if file.endswith("docx") and not file.startswith("~"):
                doc = create_doc(file_with_path)
                metadata = word_header_df(doc)
                xml_outputer(file, file_with_path,doc, metadata, database,subdir)
    global ERRORFILE
    ERRORFILE.close()

if __name__ == "__main__":
    main()
